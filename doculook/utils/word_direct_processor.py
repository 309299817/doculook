# Copyright (c) Opendatalab. All rights reserved.
"""
Word文档直接处理器
直接提取Word文档中的文本和图片，不经过PDF转换
"""

import io
import tempfile
from pathlib import Path
from typing import Dict, List, Union
from loguru import logger

try:
    import docx
    from docx import Document
except ImportError:
    logger.warning("python-docx 未安装，Word文档支持将被禁用")
    docx = None
    Document = None

try:
    from PIL import Image, ImageEnhance
except ImportError:
    logger.warning("PIL (Pillow) 未安装，图片处理功能将受限")
    Image = None
    ImageEnhance = None


def _extract_images_from_word_document(doc):
    """
    从Word文档中提取所有图片
    
    Args:
        doc: Word文档对象
        
    Returns:
        list: 图片数据列表
    """
    images = []
    try:
        # 方法1: 从文档关系中提取图片
        for rel in doc.part.rels.values():
            if hasattr(rel, 'target_ref') and "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    images.append(image_data)
                except Exception as e:
                    logger.debug(f"提取关系图片失败: {e}")
        
        # 方法2: 从document.xml中查找图片引用
        if hasattr(doc, '_element'):
            try:
                # 查找所有的图片引用
                for blip in doc._element.xpath('.//a:blip[@r:embed]'):
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id and embed_id in doc.part.rels:
                        try:
                            image_part = doc.part.rels[embed_id].target_part
                            image_data = image_part.blob
                            if image_data not in images:  # 避免重复
                                images.append(image_data)
                        except Exception as e:
                            logger.debug(f"提取blip图片失败: {e}")
            except Exception as e:
                logger.debug(f"解析document.xml失败: {e}")
                
    except Exception as e:
        logger.debug(f"提取Word图片失败: {e}")
    
    logger.info(f"从Word文档中提取到 {len(images)} 张图片")
    return images


def _enhance_image_for_ocr(pil_image):
    """
    增强图片质量以提高OCR识别效果
    
    Args:
        pil_image: PIL图片对象
        
    Returns:
        PIL.Image: 增强后的图片
    """
    if not ImageEnhance:
        return pil_image
    
    try:
        # 转换为RGB模式以确保兼容性
        if pil_image.mode not in ('RGB', 'L'):
            pil_image = pil_image.convert('RGB')
        
        # 增强对比度以提高文字清晰度
        enhancer = ImageEnhance.Contrast(pil_image)
        pil_image = enhancer.enhance(1.2)  # 增强对比度20%
        
        # 增强锐度以提高边缘清晰度
        enhancer = ImageEnhance.Sharpness(pil_image)
        pil_image = enhancer.enhance(1.1)  # 增强锐度10%
        
        return pil_image
    except Exception as e:
        logger.debug(f"图片增强失败: {e}")
        return pil_image


def process_word_document_directly(word_bytes: bytes) -> Dict:
    """
    直接处理Word文档，提取文本和图片
    
    Args:
        word_bytes: Word文档的字节数据
        
    Returns:
        dict: 包含处理结果的字典
        {
            "text_content": "提取的文本内容",
            "images": [图片字节数据列表],
            "markdown_content": "格式化的Markdown内容",
            "has_text": bool,
            "has_images": bool
        }
    """
    if Document is None:
        raise ImportError("python-docx 未安装，无法处理Word文档")
    
    try:
        # 使用BytesIO读取Word文档
        word_stream = io.BytesIO(word_bytes)
        doc = Document(word_stream)
        
        # 提取文本内容
        text_parts = []
        markdown_parts = []
        
        logger.info("开始提取Word文档文本内容...")
        
        # 按顺序处理文档中的所有元素
        logger.info("按顺序处理文档元素（段落、表格、图片）...")
        
        # 获取文档body中的所有元素，按顺序处理
        table_index = 0
        image_index = 0
        
        # 遍历文档body中的所有元素
        for element in doc.element.body:
            # 处理段落元素
            if element.tag.endswith('}p'):  # 段落元素
                # 查找对应的段落对象
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text.strip())
                            
                            # 根据段落样式生成Markdown格式
                            if paragraph.style.name.startswith('Heading'):
                                level = 1
                                if 'Heading' in paragraph.style.name:
                                    try:
                                        level = int(paragraph.style.name.replace('Heading ', ''))
                                    except:
                                        level = 1
                                markdown_parts.append(f"{'#' * level} {paragraph.text.strip()}")
                            else:
                                markdown_parts.append(paragraph.text.strip())
                        
                        # 检查段落中是否包含图片
                        for run in paragraph.runs:
                            for drawing in run._element.xpath('.//w:drawing'):
                                image_index += 1
                                # 在图片位置插入图片引用
                                image_ref = f"![图片{image_index}](images/word_image_{image_index}.png)"
                                markdown_parts.append(image_ref)
                                logger.debug(f"在段落中发现图片 {image_index}")
                        break
            
            # 处理表格元素
            elif element.tag.endswith('}tbl'):  # 表格元素
                # 查找对应的表格对象
                for table in doc.tables:
                    if table._element == element:
                        table_index += 1
                        table_markdown = []
                        table_text = []
                        
                        for row_idx, row in enumerate(table.rows):
                            row_cells = []
                            for cell in row.cells:
                                cell_text = cell.text.strip().replace('\n', ' ')
                                row_cells.append(cell_text)
                            
                            if any(cell.strip() for cell in row_cells):  # 如果行不为空
                                table_text.append(" | ".join(row_cells))
                                table_markdown.append("| " + " | ".join(row_cells) + " |")
                                
                                # 添加表格标题行分隔符
                                if row_idx == 0 and len(row_cells) > 0:
                                    separator = "| " + " | ".join(["---"] * len(row_cells)) + " |"
                                    table_markdown.append(separator)
                        
                        if table_text:
                            text_parts.append(f"\n表格 {table_index}:")
                            text_parts.extend(table_text)
                            
                            markdown_parts.append(f"\n**表格 {table_index}:**\n")
                            markdown_parts.extend(table_markdown)
                            markdown_parts.append("")
                        
                        logger.debug(f"处理表格 {table_index}")
                        break
        
        # 提取图片数据
        logger.info("提取Word文档图片...")
        document_images = _extract_images_from_word_document(doc)
        
        # 生成最终内容
        final_text = "\n\n".join([part for part in text_parts if part.strip()])
        
        # 生成Markdown内容（图片引用已经在相应位置插入）
        final_markdown = "\n\n".join([part for part in markdown_parts if part.strip()])
        
        result = {
            "text_content": final_text,
            "images": document_images,
            "markdown_content": final_markdown,
            "has_text": len(final_text.strip()) > 0,
            "has_images": len(document_images) > 0,
            "text_length": len(final_text),
            "image_count": len(document_images)
        }
        
        logger.info(f"Word文档处理完成: 文本长度={result['text_length']}, 图片数量={result['image_count']}")
        
        return result
            
    except Exception as e:
        logger.error(f"Word文档直接处理失败: {e}")
        raise


def save_word_images_for_ocr(images: List[bytes], output_dir: str) -> List[str]:
    """
    保存Word文档中的图片文件供OCR处理
    
    Args:
        images: 图片字节数据列表
        output_dir: 输出目录
        
    Returns:
        List[str]: 保存的图片文件路径列表
    """
    if not Image:
        logger.warning("PIL未安装，无法保存图片")
        return []
    
    saved_paths = []
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    for i, image_data in enumerate(images):
        try:
            # 打开图片
            image_stream = io.BytesIO(image_data)
            pil_image = Image.open(image_stream)
            
            # 增强图片质量
            pil_image = _enhance_image_for_ocr(pil_image)
            
            # 保存图片
            image_path = output_path / f"word_image_{i+1}.png"
            pil_image.save(str(image_path), format='PNG', quality=95, optimize=True)
            saved_paths.append(str(image_path))
            
            logger.info(f"保存图片: {image_path}")
            
        except Exception as e:
            logger.warning(f"保存图片 {i+1} 失败: {e}")
    
    return saved_paths


if __name__ == "__main__":
    # 测试代码
    logger.info("Word文档直接处理器加载完成")
